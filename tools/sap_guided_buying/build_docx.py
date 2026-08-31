"""Build the Word manual from the cached SAP topics.

Structure is chosen for retrieval-augmented generation rather than for reading:
every topic is its own heading-delimited section and opens with a breadcrumb, so a
chunk lifted out of the middle of the document still says what it is about.
"""

import argparse
import glob
import json
import pathlib
import re
import sys
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

HERE = pathlib.Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))

from convert import convert_topic, runs_text, tidy  # noqa: E402

MAX_HEADING = 9  # Word's built-in heading styles stop here


def add_field(paragraph, instruction):
    """Insert a Word field (TOC, PAGE, ...) as a dirty field so Word recomputes it."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, end):
        run._r.append(element)


def make_link_resolver(manifest):
    """Annotate cross-references so they survive being torn out of context."""
    titles = {t["file_path"]: t["title"] for t in manifest["topics"]}

    def resolve(anchor):
        href = (anchor.attrs.get("href") or "").strip()
        if not href or href.startswith("#"):
            return ""
        if href.startswith(("http://", "https://")):
            return f" ({href})"
        target = titles.get(href.split("#")[0])
        text = runs_text(tidy([r for r in _anchor_runs(anchor)])).strip()
        if target and target.lower() != text.lower():
            return f' (see "{target}" in this manual)'
        if target:
            return " (in this manual)"
        return ""

    return resolve


def _anchor_runs(anchor):
    from convert import collect_runs
    return collect_runs(anchor)


def style_runs(paragraph, runs, base_size=None):
    for spec in runs:
        run = paragraph.add_run(spec["text"])
        run.bold = spec["bold"]
        run.italic = spec["italic"]
        if spec["code"]:
            run.font.name = "Consolas"
        if base_size:
            run.font.size = base_size


def write_block(doc, block):
    kind = block["kind"]

    if kind == "table":
        rows = block["rows"]
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        table.autofit = True
        for r, row in enumerate(rows):
            for c, text in enumerate(row):
                cell = table.cell(r, c)
                cell.text = text
                if r == 0 and block["has_header"]:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
        doc.add_paragraph()
        return

    if kind == "caption":
        para = doc.add_paragraph()
        run = para.add_run(f"Figure: {block['text']}")
        run.italic = True
        run.font.size = Pt(9)
        return

    indent = block.get("indent", 0)

    if kind == "list_item":
        if block["ordered"]:
            # Literal numbers, not Word auto-numbering: auto-numbers restart only via
            # numbering definitions (so they would run 1..N across the whole manual) and,
            # more importantly, they do not appear in extracted plain text — which is
            # exactly what the RAG pipeline reads. Step order must survive extraction.
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5 + 0.25 * indent)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            para.add_run(f"{block['number']}. ")
        else:
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.5 + 0.25 * indent)
        style_runs(para, block["runs"])
        return

    style_name = block.get("style", "body")
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Inches(0.25 * indent)
    if style_name == "label":
        para.paragraph_format.space_before = Pt(10)
    if style_name == "code":
        para.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
    style_runs(para, block["runs"])
    if style_name == "shortdesc":
        for run in para.runs:
            run.italic = True


def build(manifest, cache_dir, out_path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    resolve = make_link_resolver(manifest)
    manual_title = manifest["title"]

    # --- title page ---------------------------------------------------------
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(manual_title)
    run.bold = True
    run.font.size = Pt(26)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("SAP Ariba Guided Buying — End User Manual")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"SAP documentation version {manifest['version']} "
        f"(build {manifest['build_no']})\nRetrieved {manifest['retrieved']}"
    )

    # --- attribution --------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Source and Attribution", level=1)
    for text in (
        f"This manual is a faithful reproduction of the SAP Help Portal guide "
        f"“{manual_title}”, version {manifest['version']}, retrieved on "
        f"{manifest['retrieved']}.",
        "The content is authored and copyrighted by SAP SE. It is reproduced here for "
        "reference only and is not original work. For the authoritative and current "
        "version, always consult the SAP Help Portal:",
        manifest["source_url"],
        "Screenshots from the original documentation have been omitted; all text, "
        "procedures, and tables are reproduced in the order SAP publishes them.",
    ):
        doc.add_paragraph(text)

    # --- table of contents --------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    toc_para = doc.add_paragraph()
    add_field(toc_para, r'TOC \o "1-3" \h \z \u')
    doc.add_paragraph(
        "If the table of contents is empty, open the document in Word and press Ctrl+A "
        "then F9 to refresh fields."
    )

    # --- body ---------------------------------------------------------------
    doc.add_page_break()
    trail = {}
    topic_count = 0
    for topic in manifest["topics"]:
        level = topic["level"]
        trail[level] = topic["title"]
        for deeper in [k for k in trail if k > level]:
            del trail[deeper]

        doc.add_heading(topic["title"], level=min(level, MAX_HEADING))
        topic_count += 1

        breadcrumb = doc.add_paragraph()
        run = breadcrumb.add_run(
            "Guided Buying (end user) > "
            + " > ".join(trail[k] for k in sorted(trail))
        )
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        payload = json.loads((cache_dir / f"{topic['loio']}.json").read_text(encoding="utf-8"))
        for block in convert_topic(payload["body"], resolve_link=resolve):
            write_block(doc, block)

    # --- page numbers -------------------------------------------------------
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer, "PAGE")

    # python-docx ships a default template whose <w:zoom> omits the required
    # w:percent attribute; Word tolerates it, strict validators and LibreOffice do not.
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")

    doc.save(out_path)
    return topic_count


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--manifest", help="path to manifest-<id>-<version>.json")
    ap.add_argument("--out", help="output .docx path")
    args = ap.parse_args()

    manifest_path = pathlib.Path(args.manifest) if args.manifest else None
    if manifest_path is None:
        candidates = sorted(glob.glob(str(HERE / "manifest-*.json")))
        if len(candidates) != 1:
            ap.error("pass --manifest; found " + (", ".join(candidates) or "none"))
        manifest_path = pathlib.Path(candidates[0])

    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    cache_dir = HERE / manifest["cache_dir"]

    if args.out:
        out_path = pathlib.Path(args.out)
    else:
        slug = re.sub(r"[^A-Za-z0-9]+", "-", manifest["title"]).strip("-")
        out_path = HERE.parents[1] / "deliverables" / f"{slug}-{manifest['version']}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    count = build(manifest, cache_dir, out_path)
    print(f"wrote {out_path} ({count} topic sections)")


if __name__ == "__main__":
    raise SystemExit(main())
